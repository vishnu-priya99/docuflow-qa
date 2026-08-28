"""DOCX parsing: one ParsedUnit per non-empty paragraph or table.

Section/heading is taken from Word's own "Heading *" paragraph styles - a
programmatic, deterministic signal set by whoever authored the document,
never inferred by an LLM.

Paragraphs and tables are walked in true document order (not python-docx's
separate .paragraphs / .tables collections, which each iterate only their
own element type and lose interleaving) - otherwise a table's section
would be whatever heading happens to be *last* in the whole file rather
than whichever section it actually sits under.
"""
from __future__ import annotations

import io

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.services.ingestion.table_rendering import cell_text
from app.services.ingestion.units import ParsedUnit


def _iter_block_items(document: Document):
    """Yields Paragraph/Table objects in the order they actually appear in
    the document body - python-docx has no built-in API for this."""
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _table_text(table: Table) -> str:
    """Renders each non-empty row as pipe-separated cells - a cell left
    genuinely blank (e.g. an unfilled sign-off field) still renders as an
    explicit "(blank)" marker rather than being silently dropped, so
    column alignment and "this was checked and found empty" both survive
    - see table_rendering.cell_text. A row with nothing in any cell is
    still skipped entirely, same as before."""
    rows_text = []
    for row in table.rows:
        raw_cells = [c.text.strip() for c in row.cells]
        if any(raw_cells):
            rows_text.append(" | ".join(cell_text(c) for c in raw_cells))
    return "\n".join(rows_text)


def parse_docx(content: bytes) -> list[ParsedUnit]:
    document = Document(io.BytesIO(content))
    units: list[ParsedUnit] = []
    current_section: str | None = None

    for index, block in enumerate(_iter_block_items(document)):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            style_name = (block.style.name if block.style else "") or ""
            if style_name.lower().startswith("heading") or style_name.lower() == "title":
                current_section = text
                continue
            units.append(
                ParsedUnit(text=text, section=current_section, paragraph_start=index, paragraph_end=index)
            )
        else:  # Table
            text = _table_text(block)
            if text:
                units.append(
                    ParsedUnit(text=text, section=current_section, paragraph_start=index, paragraph_end=index)
                )

    return units
