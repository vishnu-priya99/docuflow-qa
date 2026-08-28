"""PDF/DOCX/PPTX/TXT ingestion + semantic retrieval + citation metadata,
and multiple files coexisting correctly in one session."""
from __future__ import annotations

import pytest

from tests.conftest import ask, upload_file
from tests.factories import (
    make_docx_bytes,
    make_pdf_bytes,
    make_pdf_with_no_text_bytes,
    make_pptx_bytes,
    make_txt_bytes,
)

pytestmark = pytest.mark.asyncio


async def test_pdf_with_no_extractable_text_is_marked_failed_not_ready(client, session_id, auth_headers):
    """A scanned/image-only PDF has no text layer - pdfplumber extracts nothing.
    That must surface as a failed upload, not a silently-empty "ready" file
    that then always answers "couldn't find that information" with no clue
    why (see upload_service.NoExtractableContentError)."""
    resp = await upload_file(
        client, session_id=session_id, headers=auth_headers, filename="scanned.pdf",
        content=make_pdf_with_no_text_bytes(), content_type="application/pdf",
    )
    assert resp.status_code == 422
    body = resp.json()
    assert "extractable text" in body["detail"].lower()

    files = await client.get(f"/api/sessions/{session_id}/files", headers=auth_headers)
    uploaded = files.json()["files"][0]
    assert uploaded["status"] == "failed"
    assert uploaded["error_message"]


async def test_pdf_retrieval_preserves_page_and_section(client, session_id, auth_headers):
    pdf_bytes = make_pdf_bytes(
        [
            "FINANCIAL PERFORMANCE\nRevenue grew significantly due to new product launches in the "
            "enterprise segment this quarter.",
            "OPERATIONAL RISKS\nSupply chain disruptions affected delivery timelines during the winter.",
        ]
    )
    resp = await upload_file(
        client, session_id=session_id, headers=auth_headers, filename="annual_report.pdf", content=pdf_bytes,
        content_type="application/pdf",
    )
    assert resp.status_code == 201, resp.text
    assert resp.json()["status"] == "ready"

    result = await ask(
        client, session_id=session_id, headers=auth_headers,
        question="What does the report say about revenue and new product launches?",
    )
    assert result["question_type"] == "SEMANTIC"
    assert result["sources"], "expected at least one cited source"
    top = result["sources"][0]
    assert top["filename"] == "annual_report.pdf"
    assert top["file_type"] == "pdf"
    assert top["page_start"] == 1
    assert top["section"] == "FINANCIAL PERFORMANCE"


async def test_pdf_detects_multiple_numbered_sections_on_one_page(client, session_id, auth_headers):
    """Real documents (CAPA reports, SOPs) routinely carry more than one
    section on a single page, with numbered headings like "1.0 ...". The
    parser must scan every line (not just the first) and handle a "." in
    the heading text - see app/services/ingestion/pdf_parser.py."""
    pdf_bytes = make_pdf_bytes(
        [
            "1.0 COMPLAINT SUMMARY\n\n"
            "Fourteen product complaints were received this quarter, mostly "
            "related to guidewire advancement difficulty during insertion.\n\n"
            "2.0 ROOT CAUSE ANALYSIS\n\n"
            "The issue was traced to a supplier change in the lumen coating "
            "that increased the coefficient of friction beyond specification."
        ]
    )
    resp = await upload_file(
        client, session_id=session_id, headers=auth_headers, filename="capa.pdf",
        content=pdf_bytes, content_type="application/pdf",
    )
    assert resp.status_code == 201, resp.text

    result = await ask(
        client, session_id=session_id, headers=auth_headers,
        question="What was the root cause of the coating friction issue?",
    )
    assert result["sources"]
    assert result["sources"][0]["section"] == "2.0 ROOT CAUSE ANALYSIS"


async def test_docx_retrieval_preserves_heading_section(client, session_id, auth_headers):
    docx_bytes = make_docx_bytes(
        [
            ("Heading 1", "Introduction"),
            ("Normal", "This document explains the onboarding process for new employees."),
            ("Heading 1", "Data Privacy"),
            ("Normal", "All personal data must be encrypted at rest using AES-256 per company policy."),
        ]
    )
    resp = await upload_file(
        client, session_id=session_id, headers=auth_headers, filename="policy.docx", content=docx_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert resp.status_code == 201
    assert resp.json()["status"] == "ready"

    result = await ask(
        client, session_id=session_id, headers=auth_headers,
        question="What encryption is required for personal data at rest?",
    )
    assert result["sources"]
    top = result["sources"][0]
    assert top["filename"] == "policy.docx"
    assert top["section"] == "Data Privacy"


async def test_docx_table_gets_section_from_its_actual_position(client, session_id, auth_headers):
    """A table's section must come from wherever it actually sits in the
    document, not from document.tables' separate (position-independent)
    collection - a doc-control table before any heading must stay
    unsectioned, and a table between two headings must attribute to the
    first, not whichever heading happens to be last in the file. See
    app/services/ingestion/docx_parser.py's _iter_block_items."""
    import io as _io

    from docx import Document as _Document

    doc = _Document()
    control_table = doc.add_table(rows=1, cols=2)
    control_table.rows[0].cells[0].text = "SOP Number"
    control_table.rows[0].cells[1].text = "QA-SOP-1042-DISTINCTIVE"

    doc.add_paragraph("Purpose", style="Heading 1")
    doc.add_paragraph("This defines the non-conforming product process.")

    mid_table = doc.add_table(rows=1, cols=2)
    mid_table.rows[0].cells[0].text = "Prepared By"
    mid_table.rows[0].cells[1].text = "J. Rao Distinctive Signer"

    doc.add_paragraph("Scope", style="Heading 1")
    doc.add_paragraph("Applies to all manufacturing sites.")

    buf = _io.BytesIO()
    doc.save(buf)

    resp = await upload_file(
        client, session_id=session_id, headers=auth_headers, filename="sop.docx", content=buf.getvalue(),
    )
    assert resp.status_code == 201, resp.text

    result = await ask(
        client, session_id=session_id, headers=auth_headers,
        question="Who is J. Rao Distinctive Signer listed as (Prepared By)?",
    )
    assert result["sources"]
    # The "Prepared By" table sits between "Purpose" and "Scope" - it must
    # be attributed to "Purpose" (its actual position), not "Scope".
    assert result["sources"][0]["section"] == "Purpose"


async def test_pptx_retrieval_preserves_slide_number_and_title(client, session_id, auth_headers):
    pptx_bytes = make_pptx_bytes(
        [
            ("Welcome", "This presentation covers our roadmap for the next fiscal year."),
            ("Roadmap Q3", "We will launch the mobile app redesign and expand into European markets."),
        ]
    )
    resp = await upload_file(
        client, session_id=session_id, headers=auth_headers, filename="roadmap.pptx", content=pptx_bytes,
        content_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
    )
    assert resp.status_code == 201
    assert resp.json()["status"] == "ready"

    result = await ask(
        client, session_id=session_id, headers=auth_headers,
        question="What are the plans for the mobile app redesign?",
    )
    assert result["sources"]
    top = result["sources"][0]
    assert top["filename"] == "roadmap.pptx"
    assert top["slide_number"] == 2
    assert top["slide_title"] == "Roadmap Q3"


async def test_txt_retrieval_preserves_line_position(client, session_id, auth_headers):
    txt_bytes = make_txt_bytes("Line one is about cats.\nLine two is about dogs and puppies.\nLine three is about birds.\n")
    resp = await upload_file(
        client, session_id=session_id, headers=auth_headers, filename="notes.txt", content=txt_bytes,
        content_type="text/plain",
    )
    assert resp.status_code == 201
    assert resp.json()["status"] == "ready"

    result = await ask(client, session_id=session_id, headers=auth_headers, question="Tell me about dogs and puppies")
    assert result["sources"]
    top = result["sources"][0]
    assert top["filename"] == "notes.txt"
    assert top["line_start"] is not None


async def test_multiple_files_coexist_in_one_session(client, session_id, auth_headers):
    await upload_file(
        client, session_id=session_id, headers=auth_headers, filename="cats.txt",
        content=make_txt_bytes("Cats are independent animals that like to nap in sunny spots all day."),
    )
    await upload_file(
        client, session_id=session_id, headers=auth_headers, filename="dogs.txt",
        content=make_txt_bytes("Dogs are loyal animals that enjoy playing fetch in the park with their owners."),
    )

    files = await client.get(f"/api/sessions/{session_id}/files", headers=auth_headers)
    assert len(files.json()["files"]) == 2
    assert {f["status"] for f in files.json()["files"]} == {"ready"}

    result = await ask(client, session_id=session_id, headers=auth_headers, question="What do dogs enjoy doing?")
    assert result["sources"][0]["filename"] == "dogs.txt"
